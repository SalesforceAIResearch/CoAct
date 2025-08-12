import logging
import os
import re
import xml.etree.ElementTree as ET
import zipfile
from io import BytesIO
from typing import List, Dict, Any, Tuple

import easyocr
from PIL import Image
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import RGBColor
from odf.opendocument import load
from odf.text import P
from odf.text import Span
from rapidfuzz import fuzz
from skimage.color import deltaE_ciede2000
from skimage.color import rgb2lab

logger = logging.getLogger("desktopenv.metric.docs")


def find_default_font(config_file_path, rules):
    """Find the default font in LibreOffice Writer."""
    default_font = None
    expected_font = rules["font_name"]

    if not config_file_path:
        return 0, "Config file path is empty"

    try:
        tree = ET.parse(config_file_path)
        root = tree.getroot()

        # Define the XML namespace used in the file
        namespace = {'oor': 'http://openoffice.org/2001/registry'}

        # Search for the node containing the default font setting for LibreOffice Writer
        for elem in root.findall('.//item[@oor:path="/org.openoffice.Office.Writer/DefaultFont"]', namespace):
            for prop in elem.findall('.//prop[@oor:name="Standard"]', namespace):
                for value in prop.findall('value', namespace):
                    default_font = value.text
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error parsing config file: {str(e)}"

    if default_font == expected_font:
        return 1, f"Default font matches expected font: {expected_font}"
    else:
        return 0, f"Default font '{default_font}' does not match expected font '{expected_font}'"


def contains_page_break(docx_file, rules):
    if not docx_file:
        return 0, "Document file path is empty"

    try:
        doc = Document(docx_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening document: {str(e)}"

    try:
        expected_page_break_count = rules["page_break_count"]
    except Exception as e:
        expected_page_break_count = None

    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    page_break_count = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            br_elems = run.element.findall('.//w:br', namespaces)
            for br in br_elems:
                if br is not None and '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type' in br.attrib and \
                        br.attrib['{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type'] == 'page':
                    page_break_count += 1
    
    if expected_page_break_count is not None and page_break_count != expected_page_break_count:
        return 0, f"Page break count {page_break_count} does not match expected count {expected_page_break_count}"

    if page_break_count > 0:
        return 1, f"Document contains {page_break_count} page break(s)"
    else:
        return 0, "Document does not contain any page breaks"

def compare_docx_files(file1, file2, **options):
    ignore_blanks = options.get('ignore_blanks', True)
    ignore_case = options.get('ignore_case', False)
    ignore_order = options.get('ignore_order', False)
    content_only = options.get('content_only', False)
    fuzzy_match = options.get('fuzzy_match', False)
    delete_empty_lines = options.get('delete_empty_lines', False)

    if not file1 or not file2:
        return 0, "One or both file paths are empty"

    def get_paragraph_texts_odt(document):
        paragraphs = document.getElementsByType(P)
        paragraph_texts = []
        for paragraph in paragraphs:
            text_parts = []
            for node in paragraph.childNodes:
                if node.nodeType == node.TEXT_NODE:
                    text_parts.append(node.data)
                elif node.nodeType == node.ELEMENT_NODE and node.tagName == 'text:span':
                    # Assuming direct text content in <text:span>, for simplicity
                    for child in node.childNodes:
                        if child.nodeType == child.TEXT_NODE:
                            text_parts.append(child.data)
            paragraph_texts.append(''.join(text_parts))
        return paragraph_texts

    # Determine file types and load documents
    if file1.endswith('.docx') and file2.endswith('.docx'):
        try:
            doc1 = Document(file1)
            doc2 = Document(file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            return 0, f"Error opening DOCX files: {str(e)}"
        doc1_paragraphs = [p.text for p in doc1.paragraphs]
        doc2_paragraphs = [p.text for p in doc2.paragraphs]
        if ignore_order:
            doc1_paragraphs = sorted(doc1_paragraphs)
            doc2_paragraphs = sorted(doc2_paragraphs)
        if delete_empty_lines:
            doc1_paragraphs = [p for p in doc1_paragraphs if p.strip()]
            doc2_paragraphs = [p for p in doc2_paragraphs if p.strip()]
    elif file1.endswith('.odt') and file2.endswith('.odt'):
        try:
            doc1 = load(file1)
            doc2 = load(file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            return 0, f"Error opening ODT files: {str(e)}"
        doc1_paragraphs = get_paragraph_texts_odt(doc1)
        doc2_paragraphs = get_paragraph_texts_odt(doc2)
        if ignore_order:
            doc1_paragraphs = sorted(doc1_paragraphs)
            doc2_paragraphs = sorted(doc2_paragraphs)
        if delete_empty_lines:
            doc1_paragraphs = [p for p in doc1_paragraphs if p.strip()]
            doc2_paragraphs = [p for p in doc2_paragraphs if p.strip()]
    else:
        # Unsupported file types or mismatch
        print("Unsupported file types or mismatch between file types.")
        return 0, "Unsupported file types or mismatch between file types"

    if content_only:
        # Compare the content of the documents
        text1 = re.sub(r'\s+', ' ', '\n'.join(doc1_paragraphs)).strip()
        text2 = re.sub(r'\s+', ' ', '\n'.join(doc2_paragraphs)).strip()
        if ignore_case:
            text1, text2 = text1.lower(), text2.lower()
        similarity = fuzz.ratio(text1, text2) / 100.0
        return similarity, f"Content similarity: {similarity:.2%}"

    # Process and compare documents
    if ignore_blanks:
        text1 = re.sub(r'\s+', ' ', '\n'.join(doc1_paragraphs)).strip()
        text2 = re.sub(r'\s+', ' ', '\n'.join(doc2_paragraphs)).strip()
        if ignore_case:
            text1, text2 = text1.lower(), text2.lower()

        if fuzzy_match:
            similarity = fuzz.ratio(text1, text2) / 100.0
            return similarity, f"Fuzzy match similarity: {similarity:.2%}"
        else:
            if text1 != text2:
                return 0, "Document contents do not match (ignoring blanks)"
    else:
        if len(doc1_paragraphs) != len(doc2_paragraphs):
            print(doc1_paragraphs)
            print(doc2_paragraphs)
            print(len(doc1_paragraphs))
            print(len(doc2_paragraphs))
            return 0, f"Different number of paragraphs: {len(doc1_paragraphs)} vs {len(doc2_paragraphs)}"
        
        if fuzzy_match:
            total_similarity = 0
            if not doc1_paragraphs:
                return 1.0, "Both documents are empty"
            for p1, p2 in zip(doc1_paragraphs, doc2_paragraphs):
                if ignore_case:
                    p1, p2 = p1.lower(), p2.lower()
                total_similarity += fuzz.ratio(p1, p2) / 100.0
            
            if len(doc1_paragraphs) == 0:
                if len(doc2_paragraphs) == 0:
                    return 1.0, "Both documents are empty"
                else:
                    return 0.0, "First document is empty but second is not"

            avg_similarity = total_similarity / len(doc1_paragraphs)
            return avg_similarity, f"Average paragraph fuzzy match similarity: {avg_similarity:.2%}"
        else:
            # Compare each paragraph
            for i, (p1, p2) in enumerate(zip(doc1_paragraphs, doc2_paragraphs)):
                if ignore_case:
                    p1, p2 = p1.lower(), p2.lower()
                if p1 != p2:
                    # show the difference
                    print("=== First Paragraph ===")
                    print(f"\033[92m{repr(p1)}\033[0m")  # Green color for p1, repr() shows hidden chars
                    print("=== Second Paragraph ===")
                    print(f"\033[91m{repr(p2)}\033[0m")  # Red color for p2, repr() shows hidden chars
                    print("=" * 50)  # Clear boundary
                    return 0, f"Paragraph {i+1} does not match"

    return 1, "Documents match exactly"


def compare_init_lines(file1, file2):
    if not file1 or not file2:
        return 0, "One or both file paths are empty"

    try:
        doc1 = Document(file1)
        doc2 = Document(file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening documents: {str(e)}"

    doc1_paragraphs = [p.text for p in doc1.paragraphs]
    doc2_paragraphs = [p.text for p in doc2.paragraphs]

    # Compare each paragraph
    for i, (p1, p2) in enumerate(zip(doc1_paragraphs, doc2_paragraphs)):
        if p1 != p2:
            # print(p1)
            # print(p2)
            return 0, f"Paragraph {i+1} does not match"

    return 1, "All initial lines match"


def compare_docx_tables(docx_file1, docx_file2):
    if not docx_file1 or not docx_file2:
        return 0, "One or both file paths are empty"

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening documents: {str(e)}"

    # get list of tables in docx
    tables1 = doc1.tables
    tables2 = doc2.tables

    if len(tables1) != len(tables2):
        return 0, f"Different number of tables: {len(tables1)} vs {len(tables2)}"

    # Compare each table content
    for table_idx, (table1, table2) in enumerate(zip(tables1, tables2)):

        if len(table1.rows) != len(table2.rows) or len(table1.columns) != len(table2.columns):
            return 0, f"Table {table_idx+1} has different dimensions: {len(table1.rows)}x{len(table1.columns)} vs {len(table2.rows)}x{len(table2.columns)}"

        # Compare each cell
        for i in range(len(table1.rows)):
            for j in range(len(table1.columns)):
                if table1.cell(i, j).text.strip() != table2.cell(i, j).text.strip():
                    return 0, f"Table {table_idx+1} cell ({i+1},{j+1}) content does not match"

    return 1, "All tables match"


def compare_docx_images(docx_file1, docx_file2):
    if not docx_file1 or not docx_file2:
        return 0, "One or both file paths are empty"

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening documents: {str(e)}"

    def extract_images(doc):
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_data = rel.target_part.blob
                images.append(BytesIO(img_data))
        return images

    images1 = extract_images(doc1)
    images2 = extract_images(doc2)
    if len(images1) != len(images2):
        return 0, f"Different number of images: {len(images1)} vs {len(images2)}"
    for idx, (img1, img2) in enumerate(zip(images1, images2)):
        if Image.open(img1).tobytes() != Image.open(img2).tobytes():
            return 0, f"Image {idx+1} does not match"
    return 1, "All images match"


def compare_image_text(image_path, rule):
    if not image_path:
        return 0, "Image path is empty"
    reader = easyocr.Reader(['en'])
    result = reader.readtext(image_path)
    extracted_text = ' '.join([entry[1] for entry in result])
    
    # Log OCR results
    logger.info(f"OCR extracted texts: {[entry[1] for entry in result]}")
    logger.info(f"Combined extracted text: {extracted_text}")
    
    if rule['type'] == 'text':
        target_text = rule['text']
        match_found = target_text in extracted_text
        
        # Log matching results
        logger.info(f"Target text: '{target_text}'")
        logger.info(f"Match found: {match_found}")
        if match_found:
            logger.info("✅ Text matching successful!")
            return 1, f"Target text '{target_text}' found in image"
        else:
            logger.info("❌ Text matching failed!")
            return 0, f"Target text '{target_text}' not found in image. Extracted text: '{extracted_text}'"
    else:
        raise ValueError("Unsupported rule type")


def compare_line_spacing(docx_file1, docx_file2):
    if not docx_file1 or not docx_file2:
        return 0, "One or both file paths are empty"

    result = compare_docx_files(docx_file1, docx_file2)
    if isinstance(result, tuple):
        score, reason = result
        if score == 0:
            return 0, f"Files don't match: {reason}"
    else:
        # For backward compatibility if compare_docx_files hasn't been updated yet
        if not result:
            return 0, "Files don't match"

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening documents: {str(e)}"

    if len(doc1.paragraphs) != len(doc2.paragraphs):
        return 0, f"Different number of paragraphs: {len(doc1.paragraphs)} vs {len(doc2.paragraphs)}"

    # Compare each paragraph line spacing
    for i, (para1, para2) in enumerate(zip(doc1.paragraphs, doc2.paragraphs)):

        spacing1 = para1.paragraph_format.line_spacing
        spacing2 = para2.paragraph_format.line_spacing

        if spacing1 != spacing2:
            return 0, f"Line spacing mismatch in paragraph {i+1}: {spacing1} vs {spacing2}"

    return 1, "Line spacing matches in all paragraphs"


def compare_insert_equation(docx_file1, docx_file2):
    if not docx_file1 or not docx_file2:
        return 0, "One or both file paths are empty"

    result = compare_docx_files(docx_file1, docx_file2)
    if isinstance(result, tuple):
        score, reason = result
        if score == 0:
            return 0, f"Files don't match: {reason}"
    else:
        # For backward compatibility if compare_docx_files hasn't been updated yet
        if not result:
            return 0, "Files don't match"

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening documents: {str(e)}"

    # Compare each paragraph if it contains equation
    for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
        for run1, run2 in zip(para1.runs, para2.runs):
            if run1.element.xpath('.//w:object') and run2.element.xpath('.//w:object'):
                return 1, "Both documents contain equations"
    return 0, "No equations found in both documents"


def compare_font_names(docx_file, rules: List[Dict[str, Any]]):
    if not docx_file:
        return 0, "Document file path is empty"

    try:
        doc = Document(docx_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening document: {str(e)}"

    expected_font = rules["font_name"]

    for i, paragraph in enumerate(doc.paragraphs):
        for j, run in enumerate(paragraph.runs):
            font_name = run.font.name
            if font_name != expected_font:
                return 0, f"Font mismatch in paragraph {i+1}, run {j+1}: '{font_name}' vs expected '{expected_font}'"
    return 1, f"All text uses expected font: {expected_font}"


def compare_subscript_contains(docx_file1, docx_file2):
    if not docx_file1 or not docx_file2:
        return 0, "One or both file paths are empty"

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening documents: {str(e)}"

    for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
        for run1, run2 in zip(para1.runs, para2.runs):
            # check if two paras both contain subscript
            if run1.font.subscript and run2.font.subscript:
                return 1, "Both documents contain subscript text"
    return 0, "No matching subscript text found in both documents"


def has_page_numbers_in_footers(docx_file):
    if not docx_file:
        return 0, "Document file path is empty"

    try:
        doc = Document(docx_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening document: {str(e)}"

    for i, section in enumerate(doc.sections):
        footer = section.footer
        if footer is None:
            return 0, f"Section {i+1} has no footer"
        footer_text = footer.paragraphs[0].text if footer.paragraphs else ''
        if not any(char.isdigit() for char in footer_text):
            # if no digit in footer, then no page number
            return 0, f"Section {i+1} footer has no page numbers (no digits found)"
    return 1, "All sections have page numbers in footers"


def is_first_line_centered(docx_file):
    if not docx_file:
        return 0, "Document file path is empty"

    try:
        doc = Document(docx_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening document: {str(e)}"

    if not doc.paragraphs:
        return 0, "Document has no paragraphs"

    first_paragraph = doc.paragraphs[0]

    # check if the first line is center justified
    if first_paragraph.paragraph_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
        return 1, "First line is center-aligned"
    else:
        return 0, f"First line is not center-aligned (alignment: {first_paragraph.paragraph_format.alignment})"


def check_file_exists(directory, filename):
    if not directory or not filename:
        return 0, "Directory or filename is empty"
    file_path = os.path.join(directory, filename)
    if os.path.isfile(file_path):
        return 1, f"File exists: {file_path}"
    else:
        return 0, f"File does not exist: {file_path}"


def check_tabstops(docx_file1, docx_file2, **kwargs) -> Tuple[float, str]:
    if not docx_file1 or not docx_file2:
        return 0.0, "One or both file paths are empty"

    try:
        doc1: Document = Document(docx_file1)
        doc2: Document = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0.0, f"Error opening documents: {str(e)}"

    para1 = [p for p in doc1.paragraphs if p.text.strip()]
    para2 = [p for p in doc2.paragraphs if p.text.strip()]
    if len(para1) != len(para2): 
        return 0.0, f"Different number of non-empty paragraphs: {len(para1)} vs {len(para2)}"

    if kwargs.get('word_number_split_by_tabstop', None) is not None:
        number = kwargs['word_number_split_by_tabstop']
        index = kwargs.get('index', 0)
        for i, p1 in enumerate(para1):
            splits = p1.text.split('\t')
            if len(splits) == 0: 
                return 0.0, f"Paragraph {i+1} has no tab-separated content"
            words = list(filter(lambda x: x.strip(), re.split(r'\s', splits[index])))
            if len(words) != number: 
                return 0.0, f"Paragraph {i+1} has {len(words)} words in tab section {index}, expected {number}"

    section = doc2.sections[0]
    paragraph_width = section.page_width - section.left_margin - section.right_margin
    ignore_tabs = lambda x: x.alignment == WD_TAB_ALIGNMENT.CLEAR or (
            x.alignment == WD_TAB_ALIGNMENT.LEFT and x.position == 0)
    minus = 0.0
    for i, (p1, p2) in enumerate(zip(para1, para2)):
        # filter CLEAR tabstop and default left-0 tabstop
        tabs1 = [tst for tst in p1.paragraph_format.tab_stops if not ignore_tabs(tst)]
        tabs2 = [tst for tst in p2.paragraph_format.tab_stops if not ignore_tabs(tst)]
        if len(tabs1) != len(tabs2): 
            return 0.0, f"Paragraph {i+1} has different number of tab stops: {len(tabs1)} vs {len(tabs2)}"
        difference = 0.0
        for t1, t2 in zip(tabs1, tabs2):
            if t1.alignment != t2.alignment: 
                return 0.0, f"Tab stop alignment mismatch in paragraph {i+1}"
            difference += abs(t1.position - t2.position)
        minus += difference / paragraph_width
    score = 1 - (minus / len(para1))
    return score, f"Tab stop similarity score: {score:.2%}"


def compare_contains_image(docx_file1, docx_file2):
    if not docx_file1 or not docx_file2:
        return 0, "One or both file paths are empty"

    try:
        doc1 = Document(docx_file1)
        doc2 = Document(docx_file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening documents: {str(e)}"

    for i, (para1, para2) in enumerate(zip(doc1.paragraphs, doc2.paragraphs)):
        for j, (run1, run2) in enumerate(zip(para1.runs, para2.runs)):
            has_image1 = 'graphicData' in run1._element.xml
            has_image2 = 'graphicData' in run2._element.xml
            if has_image1 != has_image2:
                return 0, f"Image presence mismatch in paragraph {i+1}, run {j+1}: doc1 has image={has_image1}, doc2 has image={has_image2}"
    return 1, "Image presence matches in both documents"


def evaluate_colored_words_in_tables(file_path1, file_path2, **kwargs):
    if not file_path1 or not file_path2:
        return 0, "One or both file paths are empty"

    result = compare_docx_files(file_path1, file_path2)
    if isinstance(result, tuple):
        score, reason = result
        if score == 0:
            return 0, f"Files don't match: {reason}"
    else:
        # For backward compatibility if compare_docx_files hasn't been updated yet
        if not result:
            return 0, "Files don't match"

    try:
        document = Document(file_path1)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening document: {str(e)}"

    threshold = kwargs.get('threshold', 3.5)

    def _calculate_color_difference(rgb1, rgb2):
        srgb1 = [rgb1[0] / 255.0, rgb1[1] / 255.0, rgb1[2] / 255.0]
        srgb2 = [rgb2[0] / 255.0, rgb2[1] / 255.0, rgb2[2] / 255.0]
        lab1, lab2 = rgb2lab(srgb1), rgb2lab(srgb2)
        delta_e = deltaE_ciede2000(lab1, lab2)
        return delta_e

    for table_idx, table in enumerate(document.tables):
        # Iterate through rows and cells in the table
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(paragraph.runs):
                        word = run.text
                        if word:
                            first_letter = word[0].lower()

                            if first_letter in 'aeiou' and _calculate_color_difference(run.font.color.rgb,
                                                                                       RGBColor(255, 0, 0)) > threshold:
                                return 0, f"Vowel word '{word}' in table {table_idx+1}, cell ({row_idx+1},{cell_idx+1}) is not red (color difference > {threshold})"
                            elif first_letter not in 'aeiou' and _calculate_color_difference(run.font.color.rgb,
                                                                                             RGBColor(0, 0,
                                                                                                      255)) > threshold:
                                return 0, f"Non-vowel word '{word}' in table {table_idx+1}, cell ({row_idx+1},{cell_idx+1}) is not blue (color difference > {threshold})"

    return 1, "All words in tables are correctly colored (vowels=red, consonants=blue)"


def check_highlighted_words(file_path1, file_path2):
    if not file_path1 or not file_path2:
        return 0, "One or both file paths are empty"

    result = compare_docx_files(file_path1, file_path2)
    if isinstance(result, tuple):
        score, reason = result
        if score == 0:
            return 0, f"Files don't match: {reason}"
    else:
        # For backward compatibility if compare_docx_files hasn't been updated yet
        if not result:
            return 0, "Files don't match"

    doc = load(file_path1)
    highlighted = False

    for span in doc.getElementsByType(Span):
        style_name = span.getAttribute('stylename')
        if style_name:
            for automatic_style in doc.automaticstyles.childNodes:
                if automatic_style.getAttribute('name') == style_name:
                    for property in automatic_style.childNodes:
                        if property.getAttribute('backgroundcolor') == '#ffff00':
                            highlighted = True
                            break
            if highlighted:
                break

    if highlighted:
        return 0, "Document contains highlighted words (yellow background)"
    else:
        return 1, "Document does not contain any highlighted words"


def evaluate_strike_through_last_paragraph(file_path1, file_path2):
    if not file_path1 or not file_path2:
        return 0, "One or both file paths are empty"

    result = compare_docx_files(file_path1, file_path2)
    if isinstance(result, tuple):
        score, reason = result
        if score == 0:
            return 0, f"Files don't match: {reason}"
    else:
        # For backward compatibility if compare_docx_files hasn't been updated yet
        if not result:
            return 0, "Files don't match"

    try:
        document = Document(file_path1)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening document: {str(e)}"

    if not document.paragraphs:
        return 0, "Document has no paragraphs"

    # Get the last paragraph
    last_paragraph = document.paragraphs[-1]

    # Check if any run in the last paragraph has strike-through formatting
    for i, run in enumerate(last_paragraph.runs):
        if not run.font.strike:
            return 0, f"Run {i+1} in last paragraph does not have strike-through formatting"

    return 1, "All text in the last paragraph has strike-through formatting"


def evaluate_conversion(file_path):
    if not file_path:
        return 0, "File path is empty"

    try:
        document = Document(file_path)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening document: {str(e)}"

    for table_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(paragraph.runs):
                        if run.text.isupper():
                            return 0, f"Found uppercase text '{run.text}' in table {table_idx+1}, cell ({row_idx+1},{cell_idx+1})"

    for para_idx, paragraph in enumerate(document.paragraphs):
        for run_idx, run in enumerate(paragraph.runs):
            if run.text.isupper():
                return 0, f"Found uppercase text '{run.text}' in paragraph {para_idx+1}"

    return 1, "All text has been converted to lowercase"


def evaluate_spacing(file_path):
    if not file_path:
        return 0, "File path is empty"

    try:
        document = Document(file_path)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening document: {str(e)}"

    if len(document.paragraphs) < 3:
        return 0, f"Document has only {len(document.paragraphs)} paragraphs, need at least 3"

    # Check line spacing for introduction, body, and conclusion
    introduction_spacing = document.paragraphs[0].paragraph_format.line_spacing
    body_spacing = document.paragraphs[1].paragraph_format.line_spacing
    conclusion_spacing = document.paragraphs[2].paragraph_format.line_spacing
    
    if introduction_spacing == 1.0 and body_spacing == 2.0 and conclusion_spacing == 1.5:
        return 1, "Line spacing is correct: introduction=1.0, body=2.0, conclusion=1.5"
    else:
        return 0, f"Incorrect line spacing: introduction={introduction_spacing}, body={body_spacing}, conclusion={conclusion_spacing}"


def check_italic_font_size_14(path1, path2):
    if not path1 or not path2:
        return 0, "One or both file paths are empty"

    result = compare_docx_files(path1, path2)
    if isinstance(result, tuple):
        score, reason = result
        if score == 0:
            return 0, f"Files don't match: {reason}"
    else:
        # For backward compatibility if compare_docx_files hasn't been updated yet
        if not result:
            return 0, "Files don't match"

    try:
        document = Document(path1)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening document: {str(e)}"

    for para_idx, paragraph in enumerate(document.paragraphs):
        for run_idx, run in enumerate(paragraph.runs):
            if run.italic:
                # Check if font size is 14
                if run.font.size is None:
                    return 0, f"Italic text in paragraph {para_idx+1}, run {run_idx+1} has no font size set"
                elif run.font.size.pt != 14:
                    return 0, f"Italic text in paragraph {para_idx+1}, run {run_idx+1} has font size {run.font.size.pt} instead of 14"
    return 1, "All italic text has font size 14"


def evaluate_alignment(docx_path):
    if not docx_path:
        return 0, "Document path is empty"

    # Load the document
    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening document: {str(e)}"

    # Iterate through each paragraph in the document
    for para_idx, para in enumerate(doc.paragraphs):
        # Split the paragraph into individual sentences
        sentences = para.text.split('.')

        for sent_idx, sentence in enumerate(sentences):
            # Split the sentence into words
            words = sentence.strip().split()

            # Check if the sentence has at least three words
            if len(words) < 3:
                continue  # Skip sentences with less than three words

            # The first three words should be separated from the rest
            first_part = ' '.join(words[:3])
            second_part = ' '.join(words[3:])

            # Check if the sentence structure matches the pattern: first part + large space/tab + second part
            if not (first_part in sentence and second_part in sentence and sentence.find(first_part) < sentence.find(
                    second_part)):
                return 0, f"Paragraph {para_idx+1}, sentence {sent_idx+1} does not meet alignment criteria"

    return 1, "All sentences meet the alignment criteria"


def get_unique_train_ids(initial_file):  # fixed standard
    if not initial_file:
        return set(), 0

    try:
        doc = Document(initial_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return set(), 0

    train_ids = set()
    processed_lines = 0

    for para in doc.paragraphs:
        line_parts = para.text.split(',')
        if len(line_parts) == 4:
            train_id = line_parts[1].strip()
            if train_id not in train_ids:
                train_ids.add(train_id)
                processed_lines += 1

    return train_ids, processed_lines


def check_no_duplicates(initial_file, processed_file):
    if not initial_file or not processed_file:
        return 0, "One or both file paths are empty"

    # Open the document
    train_ids_ini, ini_lines = get_unique_train_ids(initial_file)

    try:
        doc_processed = Document(processed_file)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening processed file: {str(e)}"

    train_ids_pro = set()
    processed_lines = 0  # Counter for valid lines processed

    # processed
    for para_idx, para in enumerate(doc_processed.paragraphs):
        # Each line has the format: time_HH:MM:SS, train_id, station_id, platform_no
        line_parts = para.text.split(',')
        # Ensure the line has the correct format
        if len(line_parts) == 4:
            train_id = line_parts[1].strip()
            # If train_id is already in the set, it's a duplicate
            if train_id in train_ids_pro:
                return 0, f"Duplicate train_id '{train_id}' found in paragraph {para_idx+1}"
            train_ids_pro.add(train_id)
            processed_lines += 1  # Increment valid lines counter

    if train_ids_pro != train_ids_ini:
        return 0, f"Train IDs don't match between initial and processed files"
    if processed_lines != ini_lines:
        return 0, f"Number of processed lines ({processed_lines}) doesn't match initial lines ({ini_lines})"

    # No duplicates found and at least one valid line was processed
    return 1, f"No duplicates found, {processed_lines} unique train IDs processed successfully"


def compare_docx_lines(file1, file2):
    if not file1 or not file2:
        return 0, "One or both file paths are empty"

    # Read the text of the document, line by line
    try:
        doc1 = Document(file1)
        doc2 = Document(file2)
    except Exception as e:
        logger.error(f"Error: {e}")
        return 0, f"Error opening documents: {str(e)}"

    doc1_lines = [p.text.strip() for p in doc1.paragraphs if p.text.strip()]
    doc2_lines = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
    # print(doc1_lines)
    # print(doc2_lines)

    # Convert the list of lines to sets and compare
    if set(doc1_lines) == set(doc2_lines):
        return 1, f"All lines match (ignoring order), {len(doc1_lines)} unique lines"
    else:
        only_in_doc1 = set(doc1_lines) - set(doc2_lines)
        only_in_doc2 = set(doc2_lines) - set(doc1_lines)
        return 0, f"Lines don't match. {len(only_in_doc1)} lines only in doc1, {len(only_in_doc2)} lines only in doc2"


def compare_docx_files_and_ignore_new_lines(file1, file2, **options):
    ignore_blanks = options.get('ignore_blanks', True)

    if not file1 or not file2:
        return 0, "One or both file paths are empty"

    # Determine file types and load documents
    if file1.endswith('.docx') and file2.endswith('.docx'):
        try:
            doc1 = Document(file1)
            doc2 = Document(file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            return 0, f"Error opening documents: {str(e)}"

        # First, delete all the blank in paragraphs
        doc1 = [p for p in doc1.paragraphs if p.text != '']
        doc2 = [p for p in doc2.paragraphs if p.text != '']
        doc1_paragraphs = [p.text for p in doc1]
        doc2_paragraphs = [p.text for p in doc2]
    else:
        # Unsupported file types or mismatch
        print("Unsupported file types or mismatch between file types.")
        return 0, "Unsupported file types or mismatch between file types"

    # Process and compare documents
    if ignore_blanks:
        text1 = re.sub(r'\s+', ' ', '\n'.join(doc1_paragraphs)).strip()
        text2 = re.sub(r'\s+', ' ', '\n'.join(doc2_paragraphs)).strip()
        if text1 != text2:
            return 0, "Document contents don't match (ignoring blanks and newlines)"
    else:
        if len(doc1_paragraphs) != len(doc2_paragraphs):
            return 0, f"Different number of non-empty paragraphs: {len(doc1_paragraphs)} vs {len(doc2_paragraphs)}"
        # Compare each paragraph
        for i, (p1, p2) in enumerate(zip(doc1_paragraphs, doc2_paragraphs)):
            if p1 != p2:
                return 0, f"Paragraph {i+1} doesn't match"
    return 1, "Documents match (ignoring empty lines)"


# Docx file saved in the ubuntu cannot use this function to compare highlight, don't know why, deprecated
def compare_highlighted_text(file1, file2):
    if not file1 or not file2:
        return 0, "One or both file paths are empty"

    def extract_highlighted_text(file_path):
        highlighted_texts = []

        # Open the .docx file as a zip file and read the document.xml
        with zipfile.ZipFile(file_path, 'r') as docx:
            with docx.open('word/document.xml') as document_xml:
                tree = ET.parse(document_xml)
                root = tree.getroot()

        # Define the namespaces
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        }

        # Find all runs with highlight property
        for run in root.findall('.//w:r', namespaces):
            highlight = run.find('.//w:highlight', namespaces)
            if highlight is not None and highlight.get(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') != 'none':
                text = run.find('.//w:t', namespaces)
                if text is not None:
                    highlighted_texts.append(text.text)

        return highlighted_texts

    # Read the highlighted text from both documents
    doc1_highlighted = extract_highlighted_text(file1)
    doc2_highlighted = extract_highlighted_text(file2)

    # Compare the sets of highlighted text to check if they are the same
    if set(doc1_highlighted) == set(doc2_highlighted):
        return 1, f"Highlighted text matches in both documents ({len(doc1_highlighted)} highlighted sections)"
    else:
        only_in_doc1 = set(doc1_highlighted) - set(doc2_highlighted)
        only_in_doc2 = set(doc2_highlighted) - set(doc1_highlighted)
        return 0, f"Highlighted text doesn't match. {len(only_in_doc1)} only in doc1, {len(only_in_doc2)} only in doc2"


def compare_references(file1, file2, **options):
    if not file1 or not file2:
        return 0, "One or both file paths are empty"

    reference_indicator = options.get('reference_indicator', 'References')
    reference_base_result = options.get('reference_base_result', 0.5)

    # Determine file types and load documents
    if file1.endswith('.docx') and file2.endswith('.docx'):
        try:
            doc1 = Document(file1)
            doc2 = Document(file2)
        except Exception as e:
            logger.error(f"Error: {e}")
            return 0, f"Error opening documents: {str(e)}"

        doc1_paragraphs = [p.text for p in doc1.paragraphs]
        doc2_paragraphs = [p.text for p in doc2.paragraphs]
    else:
        # Unsupported file types or mismatch
        print("Unsupported file types or mismatch between file types.")
        return 0, "Unsupported file types or mismatch between file types"

    # Find the references section in the paragraphs, find the idx of the last reference_indicator in the paragraph list
    ref1_idx = doc1_paragraphs.index(reference_indicator) if reference_indicator in doc1_paragraphs else -1
    ref2_idx = doc2_paragraphs.index(reference_indicator) if reference_indicator in doc2_paragraphs else -1

    if ref1_idx == -1 and ref2_idx == -1:
        return 1, "No references section found in either document"

    if ref1_idx == -1:
        return 0, "References section missing in first document"
    if ref2_idx == -1:
        return 0, "References section missing in second document"

    # split the reference section into reference items, and remove the empty string items
    ref1 = [p for p in doc1_paragraphs[ref1_idx + 1:] if p.strip()]
    ref2 = [p for p in doc2_paragraphs[ref2_idx + 1:] if p.strip()]

    # Compare the references

    if len(ref1) != len(ref2):
        return 0, f"Different number of references: {len(ref1)} vs {len(ref2)}"

    if len(ref1) == 0:
        return 1, "Both documents have empty references sections"

    total_similarity = 0
    for r1, r2 in zip(ref1, ref2):
        # fuzzy match the references
        similarity = fuzz.ratio(r1, r2) / 100.0
        total_similarity += similarity

    result = total_similarity / len(ref1)

    epsilon = 0.01

    if result >= reference_base_result + epsilon:
        normalized_score = (result - reference_base_result) / (1 - reference_base_result)
        return normalized_score, f"References match with similarity {result:.2%} (normalized score: {normalized_score:.2%})"
    else:
        return 0, f"References similarity {result:.2%} is below threshold {reference_base_result + epsilon:.2%}"


def compare_unique_train_records(processed_file, expected_files, **kwargs):
    """
    Compares the processed file with a list of expected files containing the
    gold standard and the initial document.
    expected_files[0] should be the gold standard file.
    expected_files[1] should be the initial file.
    """
    # Debug logging to understand what we're actually receiving
    logger.info(f"DEBUG: processed_file type: {type(processed_file)}, value: {processed_file}")
    logger.info(f"DEBUG: expected_files type: {type(expected_files)}, value: {expected_files}")
    logger.info(f"DEBUG: kwargs: {kwargs}")
    
    if not processed_file or not isinstance(expected_files, list) or len(expected_files) < 2:
        logger.error("Invalid arguments: processed_file and a list of 2 expected_files are required.")
        return 0, "Invalid arguments: processed_file and a list of 2 expected_files are required"

    gold_file = expected_files[0]
    initial_file = expected_files[1]

    if not gold_file or not initial_file:
        logger.error("Gold file or initial file path is missing from expected_files list.")
        return 0, "Gold file or initial file path is missing from expected_files list"

    # Helper function to get lines and IDs from a file
    def get_lines_and_ids_from_file(file_path):
        try:
            doc = Document(file_path)
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            train_ids = [line.split(',')[1].strip() for line in lines if len(line.split(',')) == 4]
            return lines, train_ids
        except Exception as e:
            logger.error(f"Error opening or parsing file {file_path}: {e}")
            return None, None

    # Get data from all three files
    processed_lines, processed_train_ids = get_lines_and_ids_from_file(processed_file)
    if processed_lines is None: 
        return 0, f"Error reading processed file: {processed_file}"

    gold_lines, gold_train_ids = get_lines_and_ids_from_file(gold_file)
    if gold_lines is None: 
        return 0, f"Error reading gold file: {gold_file}"

    initial_lines, _ = get_lines_and_ids_from_file(initial_file)
    if initial_lines is None: 
        return 0, f"Error reading initial file: {initial_file}"
    initial_lines_set = set(initial_lines)

    # 1. Subset Check: Ensure every processed line was in the initial file
    if not set(processed_lines).issubset(initial_lines_set):
        extra_lines = set(processed_lines) - initial_lines_set
        logger.error("Processed file contains lines not present in the initial file.")
        logger.error(f"Extra lines: {extra_lines}")
        return 0, f"Processed file contains {len(extra_lines)} lines not present in the initial file"

    # 2. Uniqueness Check: Check for duplicates within the processed file
    if len(processed_train_ids) != len(set(processed_train_ids)):
        duplicate_count = len(processed_train_ids) - len(set(processed_train_ids))
        logger.error("Duplicate train_ids found in the processed file.")
        return 0, f"Found {duplicate_count} duplicate train_ids in the processed file"

    # 3. Correctness Check: Compare the set of train_ids
    if set(processed_train_ids) != set(gold_train_ids):
        missing_ids = set(gold_train_ids) - set(processed_train_ids)
        extra_ids = set(processed_train_ids) - set(gold_train_ids)
        logger.error("Set of train_ids does not match between processed file and gold file.")
        return 0, f"Train IDs mismatch: {len(missing_ids)} missing, {len(extra_ids)} extra"

    # 4. Line count check
    if len(processed_lines) != len(gold_lines):
        logger.error("Number of lines does not match between processed file and gold file.")
        return 0, f"Line count mismatch: processed has {len(processed_lines)}, gold has {len(gold_lines)}"

    return 1, f"All checks passed: {len(processed_train_ids)} unique train records processed correctly"

if __name__ == "__main__":
    image_path = "/home/ubuntu/OSWorld/cache/02ce9a50-7af2-47ed-8596-af0c230501f8/ls.png"
    print(compare_image_text(image_path, {
        "type": "text",
        "text": "ls"
      }))